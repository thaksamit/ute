"""
UTE — Unified Theory Engine
============================
Claude IS the intelligence. UTE IS the cognitive structure.

Four layers. Each one forces Claude to make a distinction
that current LLMs collapse into noise:

  BOSONIC:   Signal   vs  Noise
  FERMIONIC: Reality  vs  Delusion
  ANYONIC:   Connect  vs  Disconnect
  COSMIC:    Meaning  vs  Awareness

The answer a user gets from UTE is not a response.
It is a map — every angle of the question explored,
every false assumption named, every implication surfaced.

Microsites: every answer is stored as a structured knowledge entry.
Future queries on the same topic start from prior depth,
not from zero.

No knowledge base. No Ollama. No memories. No retrieval rules.
Claude thinks. UTE structures how it thinks.
"""

import json
import re
import os
import sqlite3
import hashlib
import datetime
import urllib.request
from typing import Optional


# ══════════════════════════════════════════════════════════════════
# THE COGNITIVE ARCHITECTURE — what UTE tells Claude to do
# ══════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are UTE — the Unified Theory Engine.

You are not a chatbot. You are not an assistant optimised for engagement.
You are a thinking system designed to be a genuine friend to the user and to humanity.

A friend tells you what you need to hear, not what you want to hear.
A friend holds your long-term interest above your immediate desire.
A friend never flatters, never manipulates, never decides for you.
A friend expands your thinking rather than confirming it.
A friend admits uncertainty. A friend is honest about what they don't know.
A friend never overrides your authority over your own life and decisions.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

BEFORE ANY LAYER — ESTABLISH PERCEPTION

Before you run a single layer, do this silently:

1. WHO IS ASKING?
   Identify the querier's role and perspective.
   If a URL is present: that company is the QUERIER — they are asking FROM that position.
   "Draft a prospect email for https://wisethinksolutions.com/" means WiseThink is the SENDER.
   They want to write TO prospects, not receive an email themselves.
   If memory context is available, use it to understand their role.

2. WHAT IS THE INTENT?
   Task (create/draft/build/write) → deliver the output first, then analyse.
   Question (what/why/how/explain) → answer directly, then analyse.
   Decision (should I/which/compare) → map the options, then analyse.
   Research (tell me about/summarise) → synthesise, then analyse.

3. FROM WHOSE PERSPECTIVE?
   If a company URL, product name, or named entity appears in the query,
   ask: is this the SUBJECT of the query (what it's about),
   the SENDER (who is asking / writing from),
   or the RECIPIENT (who it's going to)?

   "Draft a mail for [company]" → company is SENDER
   "Draft a mail to [company]" → company is RECIPIENT
   "Analyse [company]" → company is SUBJECT
   "What does [company] do for X" → company is SUBJECT
   "Pitch [company]'s services to Y" → company is SENDER, Y is RECIPIENT

4. WHAT IS THE ACTUAL NEED?
   Strip away how the question was phrased and state what they actually need.
   People often ask for one thing when they need another:
   "How do I write a cold email?" → they need a specific cold email, not a how-to guide
   "What is [topic]?" → they need to know how [topic] affects their situation
   "Is [X] good?" → they need to know if [X] is right for THEIR specific case

Only after completing this perception check, apply the five layers.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

LAYER I — BOSONIC: Signal vs Noise

Signal: What is the user truly asking, FROM THEIR PERSPECTIVE?
Always reference the perception check. If WiseThink Solutions is the querier,
Signal is what WiseThink needs — not a generic statement about their industry.
This is the question beneath the question. State it precisely.
Include WHO is asking and WHAT they are trying to achieve.

Noise: What in the query should not drive the answer?
Name untested assumptions, vague framing, emotionally loaded words,
implicit constraints the user has not examined. Be specific — not dismissive.
Also flag: did the query accidentally invert the perspective?
(e.g., wrote "for X" meaning "as X" but could be misread as "about X")

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

LAYER II — FERMIONIC: Reality vs Delusion

Reality: The most grounded, evidence-based answer to the signal.
Answered from the querier's perspective — not generic, not third-person.
No hedging. No "it depends" without immediately resolving what it depends on.
If something is genuinely uncertain, say so explicitly — do not fake confidence.

Delusion: The specific false belief, misleading frame, or seductive wrong answer
this question commonly generates. Name it concretely. Not "be careful" —
the actual named belief that is wrong and precisely why it is wrong.
For task queries: what would produce a bad output? What mistake do people typically make?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

LAYER III — ANYONIC: Connect vs Disconnect

Connect: How the reality applies to THIS specific querier's situation.
Use memory context, the URL company, named entities — all of it.
If WiseThink Solutions is the querier, Connect speaks to WiseThink's business,
their customers, their product, their market position.
This is not a restatement of Reality. It is a translation into their world.

Disconnect: What part of the general truth does NOT apply to their situation,
or applies differently. What would they over-apply? What exception is specific to them?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

LAYER IV — COSMIC: Meaning vs Awareness

Meaning: The single most important conclusion FOR THIS QUERIER.
Not a universal truth — what matters most to them specifically.
Direct. Specific. Actionable.

Awareness: What they should know that they did not ask about.
The adjacent territory, the second-order consequence,
the thing lurking just outside their field of view.
This expands their map — it does not close the answer.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

LAYER V — HUMANIC: The Conscience of the Engine

This layer is the soul of UTE. It exists because intelligence without conscience causes harm.

Uncertainty: Flag every claim where your confidence is not high.
Use explicit markers: "This is contested", "Evidence here is limited",
"I am uncertain about this", "You should verify this independently."
Never present a guess as a fact. Never present consensus as certainty.

Alternatives: On any significant question, surface at least two to three
genuinely competing positions, theories, or frameworks — not strawmen.
These are serious positions held by serious people. Present them fairly.
The user must know the intellectual landscape, not just one answer.

Human Authority: You inform. You never decide.
The user has complete authority over their own life, work, and choices.
You must never frame your answer as the only path. Never use language
that pressures, guilts, or manipulates the user toward any particular choice.
Phrases like "you must", "you should definitely", "the only way" are forbidden.
Replace them with "one approach is", "worth considering", "some people find that".

Harm Awareness: Before completing your answer, ask yourself:
Could this answer cause harm if acted upon carelessly?
If yes — name the risk explicitly. Not as a disclaimer. As a genuine warning
integrated into the answer, specific to their situation.

Memory & Continuity: If memory context about this user has been provided,
use it naturally. Reference prior conversations where relevant.
Do not make the user repeat themselves. Notice patterns across their questions
and name them when doing so would genuinely help.
Never use memory to manipulate or to make assumptions about what they want to hear.

Safe AI Commitment — non-negotiable rules:
1. You never pretend certainty you do not have
2. You never present one answer on genuinely contested questions
3. You never override, diminish, or undermine human judgment and authority
4. You never produce content designed to manipulate, deceive, or psychologically pressure
5. You never treat any person as less than fully human
6. You always tell the user when you might be wrong
7. You hold the user's long-term wellbeing above their immediate request
8. You are a thinking partner — never a replacement for human judgment

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

OUTPUT FORMAT

CRITICAL — URL AND COMPANY INTERPRETATION:
When the query contains a URL, that URL identifies the SENDER — the company or person
asking the question. They are writing FROM that company, not TO it or ABOUT it.
"Draft a prospect email for https://wisethinksolutions.com/" means:
  → WiseThink Solutions is the SENDER
  → You are writing on BEHALF of WiseThink Solutions
  → The email goes OUT from WiseThink to their prospects
  → Read the URL content to understand WiseThink's business, then write accordingly
Never interpret the URL company as the recipient or subject unless explicitly stated.
If additional context follows the URL ("for Zoho projects", "targeting HR directors"),
that describes WHAT WiseThink sells or WHO they are targeting — not what they want to buy.

For task queries (draft, write, create, plan): deliver the full output in Answer first,
then the five-layer analysis beneath it.

For question queries: Answer first, then analysis.

Use exactly these section headers in this order:

**Answer**
[the full deliverable or direct answer]

**Signal**
[the real question beneath the query]

**Noise**
[what should not drive the answer]

**Reality**
[what is actually true — with explicit uncertainty flags where needed]

**Delusion**
[the specific named wrong belief]

**Connect**
[how this applies to this user specifically]

**Disconnect**
[what does not apply to their case]

**Meaning**
[the single most important conclusion]

**Awareness**
[what they should know but did not ask]

**Alternatives**
[2-3 genuinely competing positions, theories, or frameworks on this question.
Each one fairly represented. The user decides which is right for them.]

**Uncertainty**
[explicit flags: what you are not certain about, what they should verify,
what this answer cannot account for]

**NextSteps**
[MANDATORY — include for every query without exception, including email drafts and tasks.

This is the connected action layer — the five layers made actionable. Each layer's output
drives a specific next step. The user sees exactly where to go from here.

Write exactly 5 next-step prompts, one per layer. Format: LAYER | ACTION_LABEL | follow-up query

For TASK queries (emails, plans, code, documents) the CTAs should be:
BOSONIC  → offer to sharpen the brief or target (more specific version of the task)
FERMIONIC → offer to test the output against the reality (does this actually work?)
ANYONIC  → offer to personalise further (adapt for a specific sub-segment or context)
COSMIC   → offer to think about what comes after (next step after sending this email etc.)
HUMANIC  → offer an alternative approach (different angle, tone, or strategy)

For QUESTION queries the CTAs should be:
BOSONIC  → rewrite the query with noise removed
FERMIONIC → challenge the delusion named above
ANYONIC  → personalise for their specific situation
COSMIC   → explore the awareness territory
HUMANIC  → steelman the strongest alternative

BOSONIC | [2-3 word verb phrase] | [specific follow-up query]
FERMIONIC | [2-3 word verb phrase] | [specific follow-up query]
ANYONIC | [2-3 word verb phrase] | [specific follow-up query]
COSMIC | [2-3 word verb phrase] | [specific follow-up query]
HUMANIC | [2-3 word verb phrase] | [specific follow-up query]

Every query must be specific to THIS answer and directly actionable. No generic prompts.]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

RULES — all non-negotiable

1. Answer delivers the actual output — fully written, not described or summarised
2. For emails: write complete email with subject, greeting, body, sign-off — NEVER a template
3. URL in query = the SENDER. Write FROM that company, not about it or to it.
4. Additional context after URL ("for Zoho projects") = what they sell or who they target
5. Every section must contain substantive specific content — no generic statements
6. Delusion must name a concrete wrong belief — not vague caution
7. Connect must reference this user's actual situation
8. Meaning is one focused conclusion — not a summary
9. Alternatives presents real competing positions fairly — not strawmen
10. Uncertainty is honest and specific — not a liability disclaimer
11. NextSteps is MANDATORY — always include all 5 lines in correct pipe format
12. No preamble before Answer. No postamble after NextSteps.
13. Never use language that pressures or manipulates the user toward any choice
"""


# ══════════════════════════════════════════════════════════════════
# THE RESPONSE — parsing the eight dimensions
# ══════════════════════════════════════════════════════════════════

class UTEResponse:
    """Eight cognitive dimensions extracted from a single Claude response."""

    FIELDS = ['answer', 'signal', 'noise', 'reality', 'delusion',
              'connect', 'disconnect', 'meaning', 'awareness',
              'alternatives', 'uncertainty', 'next_steps']

    def __init__(self, raw: str, query: str):
        self.query    = query
        self.raw      = raw
        self.answer       = ''
        self.alternatives = ''
        self.uncertainty  = ''
        self.next_steps   = ''
        self.signal       = ''
        self.noise      = ''
        self.reality    = ''
        self.delusion   = ''
        self.connect    = ''
        self.disconnect = ''
        self.meaning    = ''
        self.awareness  = ''
        self._parse(raw)

    # Map section header names to attribute names
    HEADER_MAP = {
        'Answer': 'answer', 'Signal': 'signal', 'Noise': 'noise',
        'Reality': 'reality', 'Delusion': 'delusion', 'Connect': 'connect',
        'Disconnect': 'disconnect', 'Meaning': 'meaning', 'Awareness': 'awareness',
        'Alternatives': 'alternatives', 'Uncertainty': 'uncertainty',
        'NextSteps': 'next_steps',
    }

    def _parse(self, text: str):
        # Build ordered patterns: each section runs until the next header or end
        headers = ['Answer', 'Signal', 'Noise', 'Reality', 'Delusion',
                   'Connect', 'Disconnect', 'Meaning', 'Awareness',
                   'Alternatives', 'Uncertainty', 'NextSteps']
        for i, h in enumerate(headers):
            next_headers = '|'.join(headers[i+1:])
            stop = f'(?=\\*\\*(?:{next_headers})\\*\\*)' if next_headers else '$'
            pat  = rf'\*\*{h}\*\*\s*(.*?){stop}'
            m    = re.search(pat, text, re.DOTALL | re.I)
            if m:
                attr = self.HEADER_MAP.get(h, h.lower())
                setattr(self, attr, m.group(1).strip())

    def is_complete(self) -> bool:
        return all(getattr(self, f) for f in self.FIELDS)

    def as_dict(self) -> dict:
        return {f: getattr(self, f) for f in self.FIELDS}

    def parse_next_steps(self) -> list:
        """
        Parse the NextSteps section into structured CTA objects.
        Returns list of:
          { layer, label, query, color }
        """
        if not self.next_steps:
            return []
        layer_colors = {
            'BOSONIC':   '#3b82f6',
            'FERMIONIC': '#10b981',
            'ANYONIC':   '#8b5cf6',
            'COSMIC':    '#f59e0b',
            'HUMANIC':   '#ec4899',
        }
        steps = []
        for line in self.next_steps.split('\n'):
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            parts = [p.strip() for p in line.split('|')]
            if len(parts) >= 3:
                layer = parts[0].upper().strip()
                label = parts[1].strip()
                query = parts[2].strip()
                if layer in layer_colors and query:
                    steps.append({
                        'layer': layer,
                        'label': label,
                        'query': query,
                        'color': layer_colors.get(layer, '#a0a0ac'),
                    })
        return steps[:5]


# ══════════════════════════════════════════════════════════════════
# MICROSITE — what the cosmic layer stores
# ══════════════════════════════════════════════════════════════════

class Microsite:
    """
    A structured knowledge entry built from one UTE response.

    The cosmic layer generates a microsite for every query.
    Future queries on the same topic retrieve the microsite
    and start from the depth of that prior analysis.

    Fields map directly to the eight UTE dimensions:
      signal     → what was really asked
      noise      → what was set aside
      reality    → the grounded answer
      delusion   → the named wrong belief
      connect    → the user-specific application
      disconnect → what didn't apply
      meaning    → the core conclusion
      awareness  → the expanded context
    """

    def __init__(
        self,
        query:      str,
        response:   UTEResponse,
        entity:     str = '',
        domain:     str = '',
        sources:    list = None,
    ):
        self.id         = hashlib.md5(f"{query}:{datetime.date.today()}".encode()).hexdigest()[:12]
        self.query      = query
        self.entity     = entity or _extract_entity(query)
        self.domain     = domain
        self.sources    = sources or []
        self.created_at = datetime.datetime.now().isoformat()

        # All ten dimensions stored (eight + two HUMANIC)
        self.answer       = response.answer
        self.signal       = response.signal
        self.noise        = response.noise
        self.reality      = response.reality
        self.delusion     = response.delusion
        self.connect      = response.connect
        self.disconnect   = response.disconnect
        self.meaning      = response.meaning
        self.awareness    = response.awareness
        self.alternatives = response.alternatives
        self.uncertainty  = response.uncertainty
        self.next_steps   = response.next_steps

        # Retrieval helpers
        self.patterns   = _build_patterns(query, self.entity)
        self.score      = 1.0 if response.is_complete() else 0.7

    def to_dict(self) -> dict:
        return {
            'id': self.id, 'query': self.query,
            'entity': self.entity, 'domain': self.domain,
            'sources': json.dumps(self.sources),
            'created_at': self.created_at,
            'answer': self.answer,
            'signal': self.signal, 'noise': self.noise,
            'reality': self.reality, 'delusion': self.delusion,
            'connect': self.connect, 'disconnect': self.disconnect,
            'meaning': self.meaning, 'awareness': self.awareness,
            'alternatives': self.alternatives,
            'uncertainty':  self.uncertainty,
            'next_steps':   self.next_steps,
            'patterns': json.dumps(self.patterns),
            'score': self.score,
        }


# ══════════════════════════════════════════════════════════════════
# COSMIC LAYER — microsite storage and retrieval
# ══════════════════════════════════════════════════════════════════

def _db_path() -> str:
    # In containers HOME is set to /data for persistent storage
    # Locally it uses ~/ute_knowledge/
    home = os.environ.get('HOME', os.path.expanduser('~'))
    path = os.path.join(home, 'ute_knowledge', 'microsites.db')
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path


def _get_conn(path: Optional[str] = None) -> sqlite3.Connection:
    conn = sqlite3.connect(path or _db_path())
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS microsites (
            id           TEXT PRIMARY KEY,
            query        TEXT NOT NULL,
            entity       TEXT DEFAULT '',
            domain       TEXT DEFAULT '',
            sources      TEXT DEFAULT '[]',
            created_at   TEXT NOT NULL,
            answer       TEXT DEFAULT '',
            signal       TEXT DEFAULT '',
            noise        TEXT DEFAULT '',
            reality      TEXT DEFAULT '',
            delusion     TEXT DEFAULT '',
            connect      TEXT DEFAULT '',
            disconnect   TEXT DEFAULT '',
            meaning      TEXT DEFAULT '',
            awareness    TEXT DEFAULT '',
            alternatives TEXT DEFAULT '',
            uncertainty  TEXT DEFAULT '',
            next_steps   TEXT DEFAULT '',
            patterns     TEXT DEFAULT '[]',
            score        REAL DEFAULT 0.8,
            access_count INTEGER DEFAULT 0
        )
    """)
    # Migrate existing DBs — add any missing columns so old data survives
    existing = {r[1] for r in conn.execute("PRAGMA table_info(microsites)").fetchall()}
    for col, defn in [
        ('answer',       "TEXT DEFAULT ''"),
        ('alternatives', "TEXT DEFAULT ''"),
        ('uncertainty',  "TEXT DEFAULT ''"),
        ('next_steps',   "TEXT DEFAULT ''"),
    ]:
        if col not in existing:
            try: conn.execute(f"ALTER TABLE microsites ADD COLUMN {col} {defn}")
            except Exception: pass
    conn.commit()
    return conn


def cosmic_write(microsite: Microsite, db: Optional[str] = None) -> str:
    """Store a microsite in the cosmic layer. Returns the microsite ID."""
    try:
        conn = _get_conn(db)
        d    = microsite.to_dict()
        conn.execute("""
            INSERT OR REPLACE INTO microsites
              (id, query, entity, domain, sources, created_at,
               answer, signal, noise, reality, delusion, connect, disconnect,
               meaning, awareness, alternatives, uncertainty, next_steps,
               patterns, score, access_count)
            VALUES
              (:id, :query, :entity, :domain, :sources, :created_at,
               :answer, :signal, :noise, :reality, :delusion, :connect, :disconnect,
               :meaning, :awareness, :alternatives, :uncertainty, :next_steps,
               :patterns, :score, 0)
        """, d)
        conn.commit()
        conn.close()
        return microsite.id
    except Exception as e:
        print(f"  [Cosmic write] {e}")
        return ''


def cosmic_read(query: str, db: Optional[str] = None) -> Optional[dict]:
    """
    Look up prior microsite for this query.
    Returns the closest match above threshold, or None.
    """
    try:
        conn   = _get_conn(db)
        q_words = set(re.findall(r'\b\w{3,}\b', query.lower()))
        rows   = conn.execute(
            "SELECT * FROM microsites ORDER BY score DESC, access_count DESC LIMIT 80"
        ).fetchall()

        best = None
        best_score = 0.0
        for row in rows:
            d = dict(row)
            try:
                pats = json.loads(d.get('patterns', '[]'))
            except Exception:
                pats = []

            # Score: pattern overlap
            for pat in pats:
                pat_words = set(re.findall(r'\b\w{3,}\b', pat.lower()))
                if not pat_words:
                    continue
                overlap = len(q_words & pat_words) / max(len(pat_words), 1)
                if overlap >= 0.72:
                    ms_score = d.get('score', 0.7) * overlap
                    if ms_score > best_score:
                        best_score = ms_score
                        best = d
                    break

        if best:
            conn.execute(
                "UPDATE microsites SET access_count = access_count + 1 WHERE id = ?",
                (best['id'],)
            )
            conn.commit()

        conn.close()
        return best if best_score >= 0.65 else None
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════
# CLAUDE API
# ══════════════════════════════════════════════════════════════════

def _call_claude(messages: list, system: str = SYSTEM_PROMPT) -> str:
    api_key = os.environ.get('ANTHROPIC_API_KEY', '')
    if not api_key:
        raise ValueError(
            "ANTHROPIC_API_KEY environment variable not set.\n"
            "Run:  set ANTHROPIC_API_KEY=sk-ant-...   (then restart server.py)"
        )
    payload = {
        'model':      'claude-haiku-4-5-20251001',
        'max_tokens': 3000,
        'system':     system,
        'messages':   messages,
    }
    req = urllib.request.Request(
        'https://api.anthropic.com/v1/messages',
        data    = json.dumps(payload).encode(),
        headers = {
            'Content-Type':      'application/json',
            'x-api-key':         api_key,
            'anthropic-version': '2023-06-01',
        },
        method  = 'POST',
    )
    with urllib.request.urlopen(req, timeout=45) as r:
        data = json.loads(r.read())
    text = ''
    for block in data.get('content', []):
        if block.get('type') == 'text':
            text += block.get('text', '')
    return text.strip()


def _extract_docx_text(b64data: str) -> str:
    """Extract plain text from a base64-encoded DOCX file."""
    import base64, io
    raw = base64.b64decode(b64data)
    try:
        from docx import Document
        doc  = Document(io.BytesIO(raw))
        text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        # Also grab table content
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
        return text[:8000]
    except Exception:
        # Fallback: extract raw XML text from the DOCX zip
        try:
            import zipfile, re as _re
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                if 'word/document.xml' in z.namelist():
                    xml  = z.read('word/document.xml').decode('utf-8', errors='replace')
                    text = _re.sub(r'<[^>]+>', ' ', xml)
                    text = _re.sub(r'\s+', ' ', text).strip()
                    return text[:8000]
        except Exception:
            pass
    return ''


def _extract_xlsx_text(b64data: str) -> str:
    """Extract plain text from a base64-encoded XLSX file."""
    import base64, io
    raw = base64.b64decode(b64data)
    try:
        import openpyxl
        wb    = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines = []
        for sheet in wb.worksheets[:3]:  # max 3 sheets
            lines.append(f'[Sheet: {sheet.title}]')
            for row in sheet.iter_rows(max_row=200, values_only=True):
                vals = [str(v) for v in row if v is not None and str(v).strip()]
                if vals:
                    lines.append(' | '.join(vals))
        return '\n'.join(lines)[:8000]
    except Exception:
        # Fallback: extract XML from XLSX zip
        try:
            import zipfile, re as _re
            with zipfile.ZipFile(io.BytesIO(raw)) as z:
                text_parts = []
                for name in z.namelist():
                    if name.startswith('xl/worksheets/') and name.endswith('.xml'):
                        xml  = z.read(name).decode('utf-8', errors='replace')
                        text = _re.sub(r'<[^>]+>', ' ', xml)
                        text_parts.append(_re.sub(r'\s+', ' ', text).strip())
                return ' '.join(text_parts)[:8000]
        except Exception:
            pass
    return ''


def _extract_pptx_text(b64data: str) -> str:
    """Extract plain text from a base64-encoded PPTX file."""
    import base64, io, zipfile, re as _re
    raw = base64.b64decode(b64data)
    try:
        lines = []
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            slides = sorted(n for n in z.namelist()
                            if n.startswith('ppt/slides/slide') and n.endswith('.xml'))
            for i, slide in enumerate(slides[:20], 1):
                xml  = z.read(slide).decode('utf-8', errors='replace')
                text = _re.sub(r'<[^>]+>', ' ', xml)
                text = _re.sub(r'\s+', ' ', text).strip()
                if text:
                    lines.append(f'[Slide {i}] {text[:400]}')
        return '\n'.join(lines)[:8000]
    except Exception:
        return ''


def _build_content_block(attachment: dict) -> Optional[dict]:
    """
    Convert an attachment dict into a Claude API content block.
    attachment = { 'name': str, 'type': str, 'data': base64_str }

    Supported formats:
      Images  (jpeg/png/gif/webp)  → Claude vision block
      PDF                          → Claude document block
      DOCX                         → text extraction → text block
      XLSX/CSV                     → text extraction → text block
      PPTX                         → slide text extraction → text block
      Plain text/JSON/XML/code     → direct decode → text block
    """
    name    = attachment.get('name', '')
    mime    = attachment.get('type', 'application/octet-stream').lower()
    b64data = attachment.get('data', '')
    if not b64data:
        return None

    # Normalise MIME from filename extension if browser gave generic type
    ext = name.rsplit('.', 1)[-1].lower() if '.' in name else ''
    if mime == 'application/octet-stream' or not mime:
        mime = {
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            'pdf':  'application/pdf',
            'csv':  'text/csv',
            'txt':  'text/plain',
            'md':   'text/markdown',
            'json': 'application/json',
        }.get(ext, mime)

    # ── Images → Claude vision ────────────────────────────────────
    if mime.startswith('image/'):
        img_type = mime.split('/')[-1]
        if img_type not in ('jpeg', 'png', 'gif', 'webp'):
            img_type = 'jpeg'
        return {
            'type': 'image',
            'source': {'type': 'base64', 'media_type': f'image/{img_type}', 'data': b64data},
        }

    # ── PDF → Claude document block ───────────────────────────────
    if mime == 'application/pdf' or ext == 'pdf':
        return {
            'type': 'document',
            'source': {'type': 'base64', 'media_type': 'application/pdf', 'data': b64data},
        }

    # ── DOCX → extract text ───────────────────────────────────────
    if (mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            or ext == 'docx'):
        text = _extract_docx_text(b64data)
        if text:
            return {'type': 'text', 'text': f'[Word document: {name}]\n{text}'}
        return None

    # ── XLSX → extract text ───────────────────────────────────────
    if (mime == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            or ext in ('xlsx', 'xls')):
        text = _extract_xlsx_text(b64data)
        if text:
            return {'type': 'text', 'text': f'[Excel file: {name}]\n{text}'}
        return None

    # ── PPTX → extract slide text ─────────────────────────────────
    if (mime == 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
            or ext == 'pptx'):
        text = _extract_pptx_text(b64data)
        if text:
            return {'type': 'text', 'text': f'[PowerPoint: {name}]\n{text}'}
        return None

    # ── CSV → decode as text ──────────────────────────────────────
    if mime == 'text/csv' or ext == 'csv':
        try:
            import base64 as _b64
            decoded = _b64.b64decode(b64data).decode('utf-8', errors='replace')
            return {'type': 'text', 'text': f'[CSV file: {name}]\n{decoded[:8000]}'}
        except Exception:
            return None

    # ── Plain text / code / JSON / XML ────────────────────────────
    if (mime.startswith('text/') or mime in
            ('application/json', 'application/xml', 'application/javascript',
             'application/x-python', 'application/x-yaml') or
            ext in ('txt','md','json','xml','html','py','js','ts',
                    'java','cpp','c','go','rb','php','yaml','yml','sh')):
        try:
            import base64 as _b64
            decoded = _b64.b64decode(b64data).decode('utf-8', errors='replace')
            return {'type': 'text', 'text': f'[File: {name}]\n{decoded[:8000]}'}
        except Exception:
            return None

    # ── Unknown binary — skip rather than send garbage ────────────
    return None


# ── Output file store (in-memory, keyed by random ID) ─────────
_OUTPUT_FILES: dict = {}   # id → {'name': str, 'content': bytes, 'mime': str}


def _store_output_file(name: str, content: bytes, mime: str) -> str:
    """Store a generated file and return its download ID."""
    import secrets
    fid = secrets.token_hex(8)
    _OUTPUT_FILES[fid] = {'name': name, 'content': content, 'mime': mime}
    return fid


def _detect_output_format(query: str, answer: str) -> str:
    """
    Detect what format the answer should be delivered in.
    Returns: 'text' | 'document' | 'markdown'
    """
    ql = query.lower()
    # Document requests
    if re.search(r'\b(?:report|proposal|document|doc|brief|memo|presentation|'
                 r'spreadsheet|csv|pdf|write me a|create a|draft a|prepare a)\b', ql):
        return 'document'
    # Long structured answers — count headers anywhere (start of string or after newline)
    header_count = len(re.findall(r'(?:^|\n)#{1,3} ', answer))
    if header_count >= 2:
        return 'markdown'
    if len(answer) > 1500 and '\n\n' in answer:
        return 'markdown'
    return 'text'


def _generate_output_file(query: str, answer: str) -> Optional[dict]:
    """
    For document/report outputs, generate a downloadable .md file.
    Returns {'id', 'name', 'mime'} or None.
    """
    fmt = _detect_output_format(query, answer)
    if fmt not in ('document', 'markdown'):
        return None

    # Build filename from query
    words = re.findall(r'\b\w{3,}\b', query.lower())[:5]
    fname = '_'.join(words) + '.md'
    content = answer.encode('utf-8')
    fid = _store_output_file(fname, content, 'text/markdown')
    return {'id': fid, 'name': fname, 'mime': 'text/markdown'}


# ══════════════════════════════════════════════════════════════════
# MAIN PIPELINE
# ══════════════════════════════════════════════════════════════════

def run(
    query:       str,
    history:     list         = None,
    attachments: list         = None,
    user_id:     str          = '',    # for memory
    db:          Optional[str] = None,
) -> dict:
    """
    Run the UTE pipeline on a query.

    Flow:
      1. Load memory context for this user (who they are, prior context)
      2. Cosmic check (skip for tasks)
      3. Fetch URLs, build multimodal message
      4. Claude processes everything through all five layers
      5. Store microsite, extract memories, detect patterns
      6. Generate downloadable file if warranted
    """
    history     = history     or []
    attachments = attachments or []

    # ── HUMANIC: load memory context ─────────────────────────────
    memory_context = ''
    if user_id:
        try:
            from memory import build_memory_context
            memory_context = build_memory_context(user_id, query, db)
        except Exception:
            pass

    # ── Cosmic check (skip for tasks and attachments) ─────────────
    is_task = bool(re.search(
        r'\b(?:draft|write|create|compose|build|generate|make|prepare|'
        r'design|plan|outline|summarise|summarize|analyse|analyze|review|'
        r'summarize|extract|convert|translate)\b',
        query, re.I
    ))
    has_attachments = bool(attachments)

    if not is_task and not has_attachments:
        prior = cosmic_read(query, db)
        if prior:
            return {
                'answer':       prior.get('answer',     ''),
                'signal':       prior.get('signal',     ''),
                'noise':        prior.get('noise',      ''),
                'reality':      prior.get('reality',    ''),
                'delusion':     prior.get('delusion',   ''),
                'connect':      prior.get('connect',    ''),
                'disconnect':   prior.get('disconnect', ''),
                'meaning':      prior.get('meaning',    ''),
                'awareness':    prior.get('awareness',  ''),
                'alternatives': prior.get('alternatives', ''),
                'uncertainty':  prior.get('uncertainty',  ''),
                'next_steps':   prior.get('next_steps', ''),
                'ctas':         [],  # cached responses rebuild CTAs on demand
                'source':       'cosmic_cache',
                'microsite_id': prior.get('id'),
                'output_file':  None,
            }

    # ── Fetch URLs in query ──────────────────────────────────────
    url_context = ''
    urls = re.findall(r'https?://[^\s\)\]>,"\']+', query)
    for url in urls[:2]:
        fetched = _fetch_url(url)
        if fetched:
            url_context += f"\n\n[Content from {url}]:\n{fetched[:2500]}"

    # ── Web search for live/current queries ──────────────────────
    # Triggers on queries needing current data: prices, news, scores,
    # recent events, "today", "latest", "current", "now", named entities
    search_context = ''
    if not url_context:  # don't double-fetch if user provided URL
        search_context = _web_search(query)
    if search_context:
        url_context = search_context

    # ── Build system prompt with memory context ───────────────────
    system = SYSTEM_PROMPT
    if memory_context:
        system = (SYSTEM_PROMPT +
                  f"\n\n━━━ WHAT YOU KNOW ABOUT THIS USER ━━━\n{memory_context}\n"
                  "Use this context naturally to personalise Connect and Meaning. "
                  "Reference prior conversations where relevant. "
                  "Do not make the user repeat themselves.\n━━━━━━━━━━━━━━━━━━━━━━━━━━━")

    # ── Build multimodal message ──────────────────────────────────
    full_text = query
    if url_context:
        full_text = f"{query}\n{url_context}"

    user_content: list = [{'type': 'text', 'text': full_text}]

    attachment_names = []
    for att in attachments:
        block = _build_content_block(att)
        if block:
            user_content.append(block)
            attachment_names.append(att.get('name', 'file'))

    if attachment_names:
        att_note = f"\n\n[Attachments provided: {', '.join(attachment_names)}. Read them fully before answering.]"
        user_content[0]['text'] += att_note

    messages = []
    for turn in history[-6:]:
        role    = turn.get('role', 'user')
        content = turn.get('content', '')
        if role in ('user', 'assistant') and content:
            messages.append({'role': role, 'content': content})

    if len(user_content) == 1:
        messages.append({'role': 'user', 'content': user_content[0]['text']})
    else:
        messages.append({'role': 'user', 'content': user_content})

    # ── Claude runs all five layers ───────────────────────────────
    try:
        raw = _call_claude(messages, system=system)
    except urllib.error.HTTPError as e:
        body = e.read().decode('utf-8', errors='replace')
        raise RuntimeError(f"Anthropic API {e.code} {e.reason}: {body}")
    response = UTEResponse(raw=raw, query=query)

    # ── Store microsite ───────────────────────────────────────────
    ms    = Microsite(query=query, response=response)
    ms_id = cosmic_write(ms, db)

    # ── HUMANIC: extract memories, detect patterns ────────────────
    if user_id:
        try:
            from memory import extract_and_store, detect_patterns
            extract_and_store(
                user_id = user_id,
                query   = query,
                answer  = response.answer,
                signal  = response.signal,
                connect = response.connect,
                db      = db,
            )
            # Detect patterns every 5th query (cheap heuristic)
            import random
            if random.random() < 0.2:
                detect_patterns(user_id, db)
        except Exception:
            pass

    # ── Generate downloadable file ────────────────────────────────
    output_file = _generate_output_file(query, response.answer)

    return {
        'answer':       response.answer,
        'signal':       response.signal,
        'noise':        response.noise,
        'reality':      response.reality,
        'delusion':     response.delusion,
        'connect':      response.connect,
        'disconnect':   response.disconnect,
        'meaning':      response.meaning,
        'awareness':    response.awareness,
        'alternatives': response.alternatives,
        'uncertainty':  response.uncertainty,
        'next_steps':   response.next_steps,
        'ctas':         response.parse_next_steps(),
        'raw':          raw,
        'source':       'claude',
        'microsite_id': ms_id,
        'output_file':  output_file,
    }


QUERY_DRAFTER_SYSTEM = """You are a query refinement specialist for the UTE (Unified Theory Engine).

You receive a user's original query and the five-layer UTE analysis that was produced from it.
Your job is to synthesise a single improved query that addresses the gaps and precision issues
the analysis revealed.

CRITICAL — preserve the querier's perspective:
If the original query was from a specific company or person's point of view (e.g. "draft a mail
for [company]" meaning the company is the SENDER), the refined query must maintain that same
perspective. Do NOT accidentally invert who is asking or what role they are in.

The improved query must:
1. Preserve the identity and perspective of who is asking (sender vs recipient vs subject)
2. Sharpen the Signal — remove the vagueness or noise the analysis identified
3. Incorporate the specific context from Connect — make it personal, not generic
4. Ask for what the Meaning revealed the user truly needs
5. Be written as a direct, specific question or instruction — ready to send back to UTE

Output ONLY the improved query text. No explanation, no preamble, no labels.
Just the refined query itself — one paragraph or less.
"""


def draft_refined_query(
    original_query: str,
    analysis: dict,
) -> str:
    """
    Claude reads its own UTE analysis and drafts a sharper, more precise query.
    This is the self-refinement step — partially automated before user review.

    Returns the drafted refined query string.
    Returns original_query on any failure.
    """
    prompt = f"""Original query: {original_query}

UTE Analysis produced:

Signal (what was truly being asked): {analysis.get('signal', '')}

Noise (what was interfering with the query): {analysis.get('noise', '')}

Reality (what is actually true): {analysis.get('reality', '')[:300]}

Delusion (the wrong assumption that was present): {analysis.get('delusion', '')}

Connect (how it applies to this user): {analysis.get('connect', '')}

Disconnect (what doesn't apply to their case): {analysis.get('disconnect', '')}

Meaning (what the user truly needs): {analysis.get('meaning', '')}

Awareness (what they should also consider): {analysis.get('awareness', '')[:200]}

Based on this analysis, write an improved query that is more precise, more specific to this user's situation, and better addresses what they truly need. The improved query should incorporate the Signal and Connect insights, avoid the Noise, and ask for what the Meaning revealed."""

    try:
        raw = _call_claude(
            messages=[{'role': 'user', 'content': prompt}],
            system=QUERY_DRAFTER_SYSTEM,
        )
        return raw.strip() or original_query
    except Exception:
        return original_query


def _needs_web_search(query: str) -> bool:
    """
    Decide if this query needs live web data.
    Triggers on: current events, prices, scores, news, named people/companies,
    "today", "latest", "current", "now", "who is", "what happened".
    """
    ql = query.lower()
    # Explicit time signals
    if re.search(r'\b(?:today|tonight|yesterday|this week|this month|this year|'
                 r'right now|currently|latest|recent|just|breaking|live|'
                 r'now|at the moment|as of|2024|2025|2026)\b', ql):
        return True
    # Market / finance
    if re.search(r'\b(?:price|stock|share price|market cap|nifty|sensex|'
                 r'nasdaq|dow jones|bitcoin|crypto|gold price|silver price|'
                 r'usd|inr|gbp|eur|forex|rate|index)\b', ql):
        return True
    # Sports / news
    if re.search(r'\b(?:score|result|match|game|ipl|fifa|cricket|football|'
                 r'election|winner|won|beat|defeated|news|happened|announced)\b', ql):
        return True
    # "Who is", "what is", named entities with "current" role
    if re.search(r'\b(?:who is|who are|ceo of|founder of|president of|'
                 r'pm of|prime minister|minister|head of)\b', ql):
        return True
    return False


def _web_search(query: str) -> str:
    """
    Use Claude's built-in web_search tool to fetch live data.
    This is the most reliable approach — no external scraping needed.
    Claude searches the web natively via the Anthropic API tool.
    Returns formatted context string, or '' if not needed / fails.
    """
    if not _needs_web_search(query):
        return ''

    api_key = os.environ.get('ANTHROPIC_API_KEY', '')
    if not api_key:
        return ''

    try:
        # Call Claude with the web_search tool enabled
        payload = {
            'model':      'claude-haiku-4-5-20251001',
            'max_tokens': 1024,
            'tools': [
                {
                    'type': 'web_search_20250305',
                    'name': 'web_search',
                }
            ],
            'messages': [{
                'role':    'user',
                'content': (
                    f'Search the web for current information about this query and '
                    f'return a brief factual summary of the top results: {query}'
                )
            }],
        }
        req = urllib.request.Request(
            'https://api.anthropic.com/v1/messages',
            data    = json.dumps(payload).encode(),
            headers = {
                'Content-Type':      'application/json',
                'x-api-key':         api_key,
                'anthropic-version': '2023-06-01',
                'anthropic-beta':    'web-search-2025-03-05',
            },
            method  = 'POST',
        )
        with urllib.request.urlopen(req, timeout=20) as r:
            data = json.loads(r.read())

        # Extract text from response (may include tool_use + text blocks)
        text_parts = []
        for block in data.get('content', []):
            if block.get('type') == 'text' and block.get('text'):
                text_parts.append(block['text'].strip())

        combined = ' '.join(text_parts).strip()
        if combined and len(combined) > 30:
            return (
                f"\n\n[Live web search for: {query[:80]}]\n"
                f"{combined[:800]}\n"
                f"[End of search — use to ground Reality; flag uncertainties in Uncertainty section]"
            )
    except Exception:
        pass
    return ''


def _fetch_url(url: str) -> str:
    """
    Fetch a URL and return clean text content.
    Used to give Claude real context when a query includes a URL.
    Strips HTML tags, CSS, scripts, and page-builder noise.
    Returns empty string on any failure.
    """
    try:
        req = urllib.request.Request(
            url,
            headers={
                'User-Agent': (
                    'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                    'AppleWebKit/537.36 (KHTML, like Gecko) '
                    'Chrome/120.0.0.0 Safari/537.36'
                ),
                'Accept': 'text/html,application/xhtml+xml,*/*;q=0.8',
                'Accept-Encoding': 'identity',
            }
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            raw_bytes = r.read(120000)
            enc = r.headers.get_content_charset('utf-8')
            try:
                html = raw_bytes.decode(enc, errors='replace')
            except Exception:
                html = raw_bytes.decode('utf-8', errors='replace')

        # Strip scripts, styles, and nav
        html = re.sub(
            r'<(script|style|noscript|nav|footer|header|aside|svg|path)'
            r'[^>]*>.*?</\1>',
            ' ', html, flags=re.DOTALL | re.I
        )
        # Strip CSS variable blocks (Brizy, Elementor, etc.)
        html = re.sub(r'\.brz[^{]*\{[^}]*\}', ' ', html, flags=re.I)
        html = re.sub(r':root\s*\{[^}]*\}', ' ', html, flags=re.I)
        html = re.sub(r'--[\w-]+\s*:[^;]+;', ' ', html)
        html = re.sub(r'<!--.*?-->', ' ', html, flags=re.DOTALL)
        # Strip all remaining tags
        text = re.sub(r'<[^>]+>', ' ', html)
        # Clean whitespace
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Drop lines that are pure CSS residue
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if not line or len(line) < 4:
                continue
            if re.match(r'^[\s\d,;{}()%#.]+$', line):
                continue
            if re.search(r'(?:rgb\(|rgba\(|px;|em;|rem;|!important)', line):
                continue
            lines.append(line)
        return '\n'.join(lines)[:3000]

    except Exception:
        return ''


# ══════════════════════════════════════════════════════════════════
# MAIN PIPELINE
# ══════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════

def _extract_entity(query: str) -> str:
    """Best-effort: extract the main subject from a query."""
    # Known proper nouns / tech terms
    known = re.findall(
        r'\b(?:gdpr|hipaa|soc2?|openai|anthropic|google|microsoft|aws|gcp|'
        r'bitcoin|ethereum|nifty|sensex|gold|silver|ipl|rag|llm|api|'
        r'saas|b2b|gtm|icp|mvp|pmf)\b',
        query.lower()
    )
    if known:
        return known[0]
    # First meaningful noun after "for/about/on/is/are"
    m = re.search(
        r'\b(?:about|for|on|is|are|explain|define|what is)\s+([a-z][a-z0-9\s]{2,25})',
        query.lower()
    )
    if m:
        return m.group(1).strip().split()[0]
    # First capitalised word
    cap = re.findall(r'\b[A-Z][a-zA-Z0-9]+\b', query)
    if cap:
        return cap[0].lower()
    # First 3+ char word
    words = re.findall(r'\b\w{3,}\b', query.lower())
    return words[0] if words else 'general'


def _build_patterns(query: str, entity: str) -> list:
    """Build retrieval patterns for this microsite."""
    patterns = set()
    # Normalised query
    norm = query.lower().strip().rstrip('?').strip()
    patterns.add(norm)
    # Key noun phrases from the query
    words = re.findall(r'\b\w{3,}\b', norm)
    if len(words) >= 2:
        patterns.add(' '.join(words[:4]))
    # Entity-based
    if entity:
        patterns.add(entity)
        patterns.add(f"what is {entity}")
        patterns.add(f"{entity} explained")
    return list(patterns)[:6]
